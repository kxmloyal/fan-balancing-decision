# DOCX导出器模块
import os
from datetime import datetime

class DocxExporter:
    def __init__(self, report_exporter):
        """
        Word (DOCX)导出器
        
        Args:
            report_exporter: 报告导出器实例
        """
        self.report_exporter = report_exporter
        self.output_folder = getattr(report_exporter, 'output_folder', 'outputs')
        self.docx_available = False
        
        # 尝试导入python-docx库
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            self.docx_available = True
            self.Document = Document
            self.Inches = Inches
            self.Pt = Pt
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        except ImportError:
            print("python-docx库不可用，Word导出功能将不可用")
            print("请安装python-docx库：pip install python-docx")
    
    def export(self, session_data, output_filename=None, task_id=None):
        """
        从会话数据导出Word (DOCX)格式报告
        
        Args:
            session_data: 会话数据，包含分析结果
            output_filename: 输出文件名
            task_id: 任务ID（可选）
            
        Returns:
            str: DOCX文件路径
            
        Raises:
            Exception: 当python-docx库不可用或导出失败时抛出异常
        """
        # 检查python-docx是否可用
        if not self.docx_available:
            raise Exception("Word导出功能依赖的python-docx库不可用，请安装python-docx库：pip install python-docx")
        
        try:
            # 如果提供了任务ID，更新任务状态
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=10, message='开始准备Word文档'
                )
            
            # 确保output_folder属性存在
            os.makedirs(self.output_folder, exist_ok=True)
            
            # 生成默认输出文件名
            if not output_filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"report_{timestamp}.docx"
            
            # 确保文件名以.docx结尾
            if not output_filename.endswith('.docx'):
                output_filename += '.docx'
            
            # 构建输出路径
            output_path = os.path.join(self.output_folder, output_filename)
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=20, message='创建Word文档'
                )
            
            # 创建Word文档
            doc = self.Document()
            
            # 填充报告信息
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            fan_model = session_data.get('fan_model', '未知')
            
            # 获取最优转速
            best_speed = "未知"
            if 'evaluation_report' in session_data and 'best_speeds' in session_data['evaluation_report']:
                best_speed = session_data['evaluation_report']['best_speeds'][0] if session_data['evaluation_report']['best_speeds'] else "未知"
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=30, message='添加报告标题和基本信息'
                )
            
            # 添加标题
            title = doc.add_heading('设备不平衡量分析报告', 0)
            title.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加扇叶型号
            subtitle = doc.add_heading(fan_model, level=1)
            subtitle.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加报告信息
            doc.add_paragraph()
            info_table = doc.add_table(rows=1, cols=3)
            info_cells = info_table.rows[0].cells
            info_cells[0].text = f"报告生成时间: {timestamp}"
            info_cells[1].text = "报告类型: Word格式分析报告"
            info_cells[2].text = f"扇叶型号: {fan_model}"
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=40, message='添加分析摘要'
                )
            
            # 添加分析摘要
            doc.add_heading('分析摘要', level=1)
            summary = doc.add_paragraph()
            summary.add_run('通过对设备在不同转速下的不平衡量数据进行统计分析，得到以下关键结论：').bold = False
            doc.add_paragraph(f"推荐最优运行转速：{best_speed}")
            doc.add_paragraph('该转速点是基于IQR（四分位距）和变异系数综合评估确定的，这两个指标反映了数据的离散程度，数值越小表示设备运行越稳定。')
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=50, message='添加统计分析结果'
                )
            
            # 添加统计分析结果
            doc.add_heading('统计分析结果', level=1)
            doc.add_paragraph(f"最优转速（综合评估）：{best_speed}（综合考虑IQR和变异系数，采用加权评分法）")
            
            if 'stats_html' in session_data:
                # 从HTML中提取表格数据并添加到Word文档
                # 这里简化处理，实际项目中可能需要更复杂的HTML解析
                doc.add_paragraph('统计分析表格数据：')
                doc.add_paragraph('（注：详细表格数据请参考HTML格式报告）')
            else:
                doc.add_paragraph('测试统计数据')
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=60, message='添加统计分析方法说明'
                )
            
            # 添加关于统计分析方法的说明
            doc.add_heading('关于统计分析方法', level=1)
            doc.add_paragraph('统计指标说明：')
            stats_bullets = [
                '平均值：反映数据的集中趋势',
                '中位数：不受极值影响的中心位置度量',
                '标准偏差：衡量数据的离散程度',
                '最小值：数据中的最小值',
                '最大值：数据中的最大值',
                'IQR（四分位距）：衡量中间50%数据的离散程度，比标准偏差更稳健',
                '变异系数(CV)：标准偏差与平均值的比值，消除了量纲影响，更适合比较不同平均水平的数据波动性'
            ]
            for bullet in stats_bullets:
                doc.add_paragraph(bullet, style='List Bullet')
            
            doc.add_paragraph('最优转速选择方法（综合评估）：')
            method_bullets = [
                '采用三级评估模型确定最优转速：',
                '1. 指标归一化处理：对每个面(P1/P2/ST)分别计算IQR和变异系数(CV)，并进行归一化处理：得分 = 1 / (1 + 指标值)',
                '2. 面内综合得分计算：对每个面的IQR得分和CV得分进行加权综合：面得分 = 0.5 × IQR得分 + 0.5 × CV得分',
                '3. 面间综合总得分计算：根据不同面的重要性进行加权综合：',
                '   - P1面权重：40%',
                '   - P2面权重：40%',
                '   - ST面权重：20%',
                '   - 总得分 = 0.4 × P1得分 + 0.4 × P2得分 + 0.2 × ST得分',
                '4. 最优转速选择：根据总得分排序，得分最高的转速为最优转速'
            ]
            for bullet in method_bullets:
                doc.add_paragraph(bullet, style='List Bullet')
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=70, message='添加优化建议和技术细节'
                )
            
            # 添加优化建议
            doc.add_heading('优化建议', level=1)
            doc.add_paragraph('基于数据分析结果，我们提出以下优化建议：')
            recommendation_bullets = [
                '首选推荐转速：建议优先选用推荐的最优运行转速，该转速下设备表现出最佳的运行稳定性',
                '次优转速选择：如果最优转速因工艺限制无法使用，可参考统计表格中其他IQR和CV值较小的转速点',
                '定期监测：建议在选定转速下建立长期监测机制，持续跟踪设备运行状态',
                '数据质量提升：为进一步提高分析准确性，建议增加每组转速下的测量样本数量',
                '多维度评估：除不平衡量外，还可结合温度、振动等其他关键指标进行综合评估'
            ]
            for i, bullet in enumerate(recommendation_bullets, 1):
                doc.add_paragraph(f"{i}. {bullet}")
            
            # 添加技术细节说明
            doc.add_heading('技术细节说明', level=1)
            doc.add_paragraph('关于数据处理和分析方法的技术说明：')
            tech_bullets = [
                '所有数据均经过预处理，去除明显异常值以保证分析结果的可靠性',
                'IQR和CV作为互补指标，分别从绝对和相对角度评估数据稳定性',
                '加权评分法考虑了不同测量面的重要性差异，更符合实际工程情况',
                '图表采用箱线图和小提琴图形式，能够直观展示数据分布特征和离群点情况',
                '分析结果受测量精度和样本数量影响，建议结合实际情况进行判断'
            ]
            for bullet in tech_bullets:
                doc.add_paragraph(bullet, style='List Bullet')
            
            # 添加使用说明
            doc.add_heading('使用说明', level=1)
            doc.add_paragraph('详细的分析数据和图表请参考HTML格式报告，包括：')
            usage_bullets = [
                '各转速点的统计分析结果',
                '不同面的不平衡量图表（PNG和交互式HTML格式）'
            ]
            for bullet in usage_bullets:
                doc.add_paragraph(bullet, style='List Bullet')
            
            # 添加注意事项
            doc.add_heading('注意事项', level=1)
            note_bullets = [
                'IQR（四分位距）和变异系数反映了数据的离散程度，数值越小表示数据越稳定',
                '建议关注这些指标较小的转速点，这些点通常代表设备运行较稳定的状态',
                '如需进一步分析，请结合设备的实际运行情况进行综合判断',
                '本报告提供的最优转速建议仅供参考，实际应用中还需考虑工艺要求和其他工程因素',
                '报告中的图表和数据可下载保存，供后续分析和汇报使用'
            ]
            for bullet in note_bullets:
                doc.add_paragraph(bullet, style='List Bullet')
            
            # 添加页脚
            for section in doc.sections:
                footer = section.footer
                footer_paragraph = footer.paragraphs[0]
                footer_paragraph.text = '本报告由扇叶平衡补土转速评估工具自动生成'
                footer_paragraph.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=85, message='准备保存Word文档'
                )
            
            # 保存Word文档
            doc.save(output_path)
            
            # 更新任务进度
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'in_progress', progress=90, message='Word文档保存完成'
                )
            
            # 添加到导出历史
            export_info = {
                'type': 'docx',
                'filename': os.path.basename(output_path),
                'path': output_path,
                'fan_model': session_data.get('fan_model', '未知')
            }
            self.report_exporter.history_manager.add_record(export_info)
            
            # 更新任务状态为完成
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'completed', progress=100, message='Word报告导出完成', result=output_path
                )
            
            return output_path
        except Exception as e:
            # 更新任务状态为失败
            if task_id:
                self.report_exporter.task_manager.update_task_status(
                    task_id, 'failed', progress=0, message='Word报告导出失败', error=str(e)
                )
            print(f"导出Word报告失败: {str(e)}")
            raise
